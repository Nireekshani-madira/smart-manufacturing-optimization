# scripts/generate_report.py

import os
import joblib
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.metrics import classification_report, confusion_matrix
from docx import Document
from docx.shared import Inches

# Create report folder if it doesn't exist
if not os.path.exists("report"):
    os.makedirs("report")

# Paths
MODEL_PATH = "models/xgb_model.pkl"
PREDICTIONS_PATH = "outputs/predictions.csv"
REPORT_PATH = "report/report_generated.docx"

# Load processed data and predictions
df = pd.read_csv(PREDICTIONS_PATH)

# Load trained model
model = joblib.load(MODEL_PATH)

# Identify features used in model
cols_to_drop = [col for col in ['failure', 'predicted_failure', 'timestamp'] if col in df.columns]
features = df.drop(columns=cols_to_drop).columns

# ========================
# Confusion Matrix & Report
# ========================
cm = confusion_matrix(df['failure'], df['predicted_failure'])
plt.figure(figsize=(6,5))
sns.heatmap(cm, annot=True, fmt="d", cmap='Blues')
plt.title("Confusion Matrix")
plt.xlabel("Predicted")
plt.ylabel("Actual")
plt.tight_layout()
plt.savefig("report/confusion_matrix.png")
plt.close()

classification_rep = classification_report(df['failure'], df['predicted_failure'])
print(classification_rep)

# ========================
# Feature Importance
# ========================
importance = model.feature_importances_
importance_df = pd.DataFrame({'Feature': features, 'Importance': importance}).sort_values(by='Importance', ascending=False)

plt.figure(figsize=(8,6))
plt.barh(importance_df['Feature'], importance_df['Importance'], color='skyblue')
plt.xlabel("Importance")
plt.title("Feature Importance")
plt.gca().invert_yaxis()
plt.tight_layout()
plt.savefig("report/feature_importance.png")
plt.close()

# ========================
# Create Word Report
# ========================
doc = Document()
doc.add_heading("Smart Manufacturing Process Optimization Report", level=0)

# Confusion matrix
doc.add_heading("Model Evaluation", level=1)
doc.add_picture("report/confusion_matrix.png", width=Inches(5))
doc.add_heading("Classification Report", level=2)
doc.add_paragraph(classification_rep)

# Feature importance
doc.add_heading("Feature Importance", level=1)
doc.add_picture("report/feature_importance.png", width=Inches(5))
doc.add_heading("Top Features Table", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light List Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Feature'
hdr_cells[1].text = 'Importance'
for _, row in importance_df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(row['Feature'])
    row_cells[1].text = f"{row['Importance']:.4f}"

# Save Word report
doc.save(REPORT_PATH)
print(f"Report generated at: {REPORT_PATH}")
